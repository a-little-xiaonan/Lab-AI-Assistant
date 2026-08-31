"""文档解析：格式分发 + 四级兜底链（对齐设计文档 §4.1 与 Phase 1-04）。

格式分发表：
  PDF    → pymupdf      逐页提取块，跨页段落合并，页作为元数据
  DOCX   → python-docx  按 Heading 样式分节，无标题样式退化为按段落
  MD/TXT → 直读         优先 UTF-8，失败回退 GBK
  PPTX   → unstructured 按 slide（page_number）分组，不跨 slide
  XLSX   → openpyxl     说明：unstructured 0.27 把整个 sheet 拍平成单块文本、
                        无行边界，产不出规格要求的"~50 行/块 + sheet_name/row_range"，
                        故 XLSX 用 openpyxl 直接读行结构（openpyxl 是 unstructured
                        的自带依赖，不新增负担）。PPTX 仍走 unstructured。
  HTML   → unstructured 按元素
  未知格式 → 四级兜底链：L1 专用解析器 → L2 unstructured 自动分区 → L3 文本试探
            （UTF-8/GBK 解码）→ L4 优雅拒绝（标记 failed，不硬撑）
"""
from __future__ import annotations

import hashlib
import logging
import re
from pathlib import Path
from typing import Callable

from app.core.cleaner import clean
from app.core.models import RawElement

logger = logging.getLogger(__name__)

# 扩展名 → 格式（未知扩展名走兜底链）
SUPPORTED: dict[str, str] = {
    ".pdf": "pdf",
    ".docx": "docx",
    ".md": "md",
    ".markdown": "md",
    ".txt": "txt",
    ".pptx": "pptx",
    ".xlsx": "xlsx",
    ".html": "html",
    ".htm": "html",
}

MAX_UPLOAD_MB = 50

# 纯页码行（PDF 页脚最常见形态，整块只有数字/页码字样时剔除）
_PAGE_NUMBER_ONLY = re.compile(r"^\s*(第?\s*\d+\s*页?|\d+\s*/\s*\d+|\d+)\s*$")
# 句子结尾标点：跨页合并判定
_SENTENCE_END = "。！？!?；;"


class DocumentParseError(Exception):
    """解析失败（格式识别失败/解码失败/解析器内部错误）。"""


class UnsupportedFormatError(DocumentParseError):
    """L4 优雅拒绝：格式不受支持。"""


def detect_format(path: Path) -> str | None:
    return SUPPORTED.get(Path(path).suffix.lower())


def _read_text(path: Path) -> str:
    """UTF-8 优先，GBK 回退（Windows 老文本常见）。"""
    data = path.read_bytes()
    for enc in ("utf-8", "gbk"):
        try:
            return data.decode(enc)
        except UnicodeDecodeError:
            continue
    raise DocumentParseError(f"文件不是可解码的文本（非 UTF-8/GBK）：{path.name}")


def _split_paragraphs(text: str) -> list[RawElement]:
    return [RawElement(text=p) for p in re.split(r"\n\s*\n", text) if p.strip()]


# ---------- 各格式解析 ----------

def _extract_pdf(path: Path) -> list[RawElement]:
    import pymupdf

    elements: list[RawElement] = []
    pending: tuple[str, int] | None = None  # (跨页未完成段落文本, 起始页)
    with pymupdf.open(path) as doc:
        if doc.page_count == 0:
            return []
        for pno in range(doc.page_count):
            page = doc.load_page(pno)
            blocks = [
                (b[4].strip(), pno)
                for b in page.get_text("blocks")
                if b[6] == 0 and b[4].strip() and not _PAGE_NUMBER_ONLY.match(b[4].strip())
            ]
            for i, (text, _) in enumerate(blocks):
                if pending is not None:
                    text = pending[0] + "\n" + text
                    start_page = pending[1]
                    pending = None
                else:
                    start_page = pno
                # 页内最后一块且未以句末标点结尾 → 视为跨页段落，与下一页首块合并
                if i == len(blocks) - 1 and text[-1] not in _SENTENCE_END:
                    pending = (text, start_page)
                else:
                    elements.append(RawElement(text=text, page=start_page))
        if pending is not None:
            elements.append(RawElement(text=pending[0], page=pending[1]))
    if not elements:
        logger.warning("PDF 无文本层（扫描件？）：%s（OCR 不在 MVP 范围，跳过）", path.name)
    return elements


def _extract_docx(path: Path) -> list[RawElement]:
    from docx import Document

    doc = Document(path)
    elements: list[RawElement] = []
    section: list[str] = []
    has_heading = any(
        p.text.strip() and p.style and p.style.name.startswith("Heading") for p in doc.paragraphs
    )

    def flush():
        if section:
            elements.append(RawElement(text="\n".join(section)))
            section.clear()

    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            continue
        is_heading = p.style is not None and p.style.name.startswith("Heading")
        if is_heading:
            flush()
            section.append(text)  # 标题并入该节，检索时带上下文
        else:
            section.append(text)
    flush()

    # 表格：每行以 " | " 连接单元格
    for table in doc.tables:
        rows = [" | ".join(cell.text.strip() for cell in row.cells) for row in table.rows]
        rows = [r for r in rows if r.strip()]
        if rows:
            elements.append(RawElement(text="\n".join(rows)))
    return elements


def _extract_md(path: Path) -> list[RawElement]:
    text = _read_text(path)
    elements: list[RawElement] = []
    section: list[str] = []
    has_heading = False

    def flush():
        if section:
            elements.append(RawElement(text="\n".join(section)))
            section.clear()

    for line in text.split("\n"):
        if re.match(r"^#{1,6}\s+\S", line):
            flush()
            has_heading = True
            section.append(line)
        else:
            section.append(line)
    flush()
    # 无标题的 MD 退化按段落
    return elements if has_heading else _split_paragraphs(text)


def _extract_txt(path: Path) -> list[RawElement]:
    return _split_paragraphs(_read_text(path))


def _extract_pptx(path: Path) -> list[RawElement]:
    from unstructured.partition.pptx import partition_pptx

    elements = partition_pptx(filename=str(path))
    slides: dict[int, list[str]] = {}
    for el in elements:
        text = (el.text or "").strip()
        if not text:
            continue
        slide = getattr(el.metadata, "page_number", None) or 1
        slides.setdefault(int(slide), []).append(text)
    return [
        RawElement(text="\n".join(slides[slide]), slide_number=slide)
        for slide in sorted(slides)
    ]


def _extract_xlsx(path: Path) -> list[RawElement]:
    """openpyxl 直读行结构（见模块 docstring 的说明）。每 ~50 行一块，块间重复表头。"""
    import openpyxl

    BLOCK_ROWS = 50
    wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
    elements: list[RawElement] = []
    try:
        for ws in wb.worksheets:
            rows: list[str] = []
            for row in ws.iter_rows(values_only=True):
                cells = [str(c).strip() for c in row if c is not None and str(c).strip()]
                if cells:
                    rows.append(" | ".join(cells))
            if not rows:
                continue
            header = rows[0]
            for start in range(0, len(rows), BLOCK_ROWS):
                block = rows[start : start + BLOCK_ROWS]
                parts = ([header] if start > 0 else []) + block
                elements.append(
                    RawElement(
                        text="\n".join(parts),
                        sheet_name=ws.title,
                        row_range=f"{start + 1}-{start + len(block)}",
                    )
                )
    finally:
        wb.close()
    return elements


def _extract_html(path: Path) -> list[RawElement]:
    from unstructured.partition.html import partition_html

    elements = partition_html(filename=str(path))
    return [RawElement(text=el.text.strip()) for el in elements if el.text and el.text.strip()]


_PARSERS: dict[str, Callable[[Path], list[RawElement]]] = {
    "pdf": _extract_pdf,
    "docx": _extract_docx,
    "md": _extract_md,
    "txt": _extract_txt,
    "pptx": _extract_pptx,
    "xlsx": _extract_xlsx,
    "html": _extract_html,
}


# ---------- 兜底链 ----------

def _fallback_chain(path: Path) -> tuple[str, list[RawElement]]:
    """L2 unstructured 自动分区 → L3 文本试探 → L4 优雅拒绝。"""
    # L2
    try:
        from unstructured.partition.auto import partition

        elements = partition(filename=str(path))
        texts = [el.text.strip() for el in elements if el.text and el.text.strip()]
        if texts:
            return "auto", [RawElement(text=t) for t in texts]
    except Exception as exc:
        logger.info("L2 unstructured 解析失败（%s）：%s", path.name, exc)
    # L3
    try:
        text = _read_text(path)
        if text.strip():
            return "txt", _split_paragraphs(text)
    except DocumentParseError as exc:
        logger.info("L3 文本解码失败（%s）：%s", path.name, exc)
    # L4
    raise UnsupportedFormatError(
        f"暂不支持该格式（{path.suffix or '未知扩展名'}），请转为 PDF / DOCX / MD / TXT 后上传"
    )


# ---------- 对外入口 ----------

def load(path: Path) -> tuple[str, list[RawElement]]:
    """解析文档 → (格式, 结构化元素列表)。失败抛 DocumentParseError（含 UnsupportedFormatError）。"""
    path = Path(path)
    if path.stat().st_size > MAX_UPLOAD_MB * 1024 * 1024:
        raise DocumentParseError(f"文件超过 {MAX_UPLOAD_MB}MB 大小限制")
    fmt = detect_format(path)
    if fmt and fmt in _PARSERS:
        try:
            elements = _PARSERS[fmt](path)
        except UnsupportedFormatError:
            raise
        except DocumentParseError:
            raise
        except Exception as exc:
            logger.exception("解析 %s 失败", path.name)
            raise DocumentParseError(f"解析失败：{exc}") from exc
        return fmt, [RawElement(text=clean(el.text), page=el.page, slide_number=el.slide_number,
                                sheet_name=el.sheet_name, row_range=el.row_range) for el in elements]
    fmt, elements = _fallback_chain(path)
    return fmt, [RawElement(text=clean(el.text), page=el.page, slide_number=el.slide_number,
                            sheet_name=el.sheet_name, row_range=el.row_range) for el in elements]


def file_hash(path: Path) -> str:
    """内容 hash：重复上传去重用（sha256 前 16 位足够）。"""
    h = hashlib.sha256()
    with open(path, "rb") as f:
        for block in iter(lambda: f.read(65536), b""):
            h.update(block)
    return h.hexdigest()[:16]
